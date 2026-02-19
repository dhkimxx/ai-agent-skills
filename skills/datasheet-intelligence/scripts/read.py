#!/usr/bin/env python3
"""파일 포맷을 자동 감지하여 텍스트를 추출한다.

지원 포맷:
  - PDF: 특정 페이지 범위의 원본 텍스트 (--structured: docling 구조 파싱)
  - DOCX: 구조 파싱 (단락/테이블) (--structured: docling)
  - XLSX: 시트 데이터 (--structured: docling)
"""

import argparse
import json
import os
import sys
import tempfile
from pathlib import Path


def print_next_step(tip: str, command: str | None = None) -> None:
    """에러 직후 다음 액션을 동일 포맷(Tip/Try)으로 안내한다."""
    print(f"Tip: {tip}", file=sys.stderr)
    if command:
        print(f"Try: {command}", file=sys.stderr)


# ---------------------------------------------------------------------------
# 공통: 구조적 파싱 (docling)
# ---------------------------------------------------------------------------


def read_structured(path: Path, pages: list[int] | None = None) -> None:
    """docling으로 파일을 구조적 파싱하여 마크다운을 출력한다."""
    try:
        from docling.document_converter import DocumentConverter
    except ImportError:
        print("Error: --structured 모드는 docling이 필요합니다.", file=sys.stderr)
        print_next_step(
            "docling extra를 포함해 structured 모드로 다시 실행하세요.",
            "uv run --project skills/datasheet-intelligence --with docling "
            "skills/datasheet-intelligence/scripts/read.py <file> --structured",
        )
        sys.exit(1)

    tmp_path = None
    target_path = path

    try:
        # PDF인 경우 페이지만 분리해서 처리 (속도 최적화)
        if path.suffix.lower() == ".pdf" and pages:
            print(
                f"[structured] {len(pages)}페이지를 docling으로 파싱 중...",
                file=sys.stderr,
            )
            target_path = Path(_split_pdf_pages(path, pages))
            tmp_path = target_path
        else:
            print(
                f"[structured] {path.name} 전체를 docling으로 파싱 중...",
                file=sys.stderr,
            )

        converter = DocumentConverter()
        result = converter.convert(target_path)
        md = result.document.export_to_markdown()

        print(md)

    except Exception as e:
        print(f"Error: 구조적 파싱 실패 — {e}", file=sys.stderr)
        base_command = (
            "uv run --project skills/datasheet-intelligence "
            "skills/datasheet-intelligence/scripts/read.py "
            f'"{path}"'
        )
        if pages:
            page_args = ",".join(str(page) for page in pages)
            base_command += f" --pages {page_args}"
        print_next_step(
            "먼저 기본 모드로 같은 범위를 확인한 뒤, 필요한 경우에만 --structured를 사용하세요.",
            base_command,
        )
        sys.exit(1)
    finally:
        if tmp_path and tmp_path.exists():
            os.unlink(tmp_path)


# ---------------------------------------------------------------------------
# PDF — 기본 모드 (pypdfium2, 빠름)
# ---------------------------------------------------------------------------


def _parse_page_range(range_str: str) -> list[int]:
    pages: list[int] = []
    for part in range_str.split(","):
        if "-" in part:
            start, end = map(int, part.split("-"))
            pages.extend(range(start, end + 1))
        else:
            pages.append(int(part))
    return sorted(set(pages))


def _split_pdf_pages(src_path: Path, pages: list[int]) -> str:
    """지정 페이지만 임시 PDF로 분리한다 (pypdfium2)."""
    import pypdfium2 as pdfium

    src = pdfium.PdfDocument(src_path)
    total = len(src)
    valid = sorted(p - 1 for p in pages if 1 <= p <= total)  # 0부터 시작하는 인덱스로 변환

    if not valid:
        raise ValueError(f"유효한 페이지 없음. 전체 {total}페이지.")

    new_doc = pdfium.PdfDocument.new()
    new_doc.import_pages(src, pages=valid)

    tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    new_doc.save(tmp)
    tmp.close()
    return tmp.name


def read_pdf(path: Path, pages: list[int]) -> None:
    import pypdfium2 as pdfium

    pdf = pdfium.PdfDocument(path)
    total = len(pdf)
    valid = [p for p in pages if 1 <= p <= total]

    if not valid:
        print(f"Error: 유효한 페이지 없음. 전체 {total}페이지.", file=sys.stderr)
        print_next_step(
            "먼저 전체 페이지 수를 확인한 뒤 올바른 범위를 지정하세요.",
            "uv run --project skills/datasheet-intelligence "
            f'skills/datasheet-intelligence/scripts/read.py "{path}"',
        )
        return

    for num in valid:
        page = pdf.get_page(num - 1)
        tp = page.get_textpage()
        print(f"\n=== PAGE {num} ===")
        print(tp.get_text_range())
        tp.close()
        page.close()


# ---------------------------------------------------------------------------
# DOCX — 기본 모드 (python-docx, 빠름)
# ---------------------------------------------------------------------------


def read_docx(path: Path) -> None:
    from docx import Document

    doc = Document(path)
    for i, para in enumerate(doc.paragraphs, 1):
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else "Normal"
        print(
            json.dumps({"index": i, "style": style, "text": text}, ensure_ascii=False)
        )


# ---------------------------------------------------------------------------
# XLSX — 기본 모드 (openpyxl, 빠름)
# ---------------------------------------------------------------------------


def read_xlsx(path: Path) -> None:
    from openpyxl import load_workbook

    wb = load_workbook(path, read_only=True, data_only=True)
    for name in wb.sheetnames:
        print(f"\n=== Sheet: {name} ===")
        for row in wb[name].iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(c.strip() for c in cells):
                print("\t".join(cells))
    wb.close()


# ---------------------------------------------------------------------------
# 메인 진입점
# ---------------------------------------------------------------------------


def main():
    parser = argparse.ArgumentParser(
        description="파일에서 텍스트를 추출한다 (PDF/DOCX/XLSX)."
    )
    parser.add_argument("file_path", type=Path, help="대상 파일 경로")
    parser.add_argument(
        "--pages", help="PDF 페이지 범위 (예: '10', '10-20', '10,15,20')"
    )
    parser.add_argument(
        "--structured",
        action="store_true",
        help="docling으로 구조적 파싱 (테이블/헤딩 구조 보존, 느리지만 정확)",
    )
    args = parser.parse_args()

    if not args.file_path.exists():
        print(f"Error: 파일 없음 — {args.file_path}", file=sys.stderr)
        print_next_step(
            "파일 경로를 다시 확인하세요.",
            f'ls -l "{args.file_path}"',
        )
        sys.exit(1)

    # --structured 모드 공통 처리 (포맷 무관)
    if args.structured:
        # PDF인 경우만 pages 인자 사용, 나머지는 전체 파싱
        pages = _parse_page_range(args.pages) if args.pages else None
        read_structured(args.file_path, pages)
        return

    # 기본 모드 처리
    suffix = args.file_path.suffix.lower()

    if suffix == ".pdf":
        if not args.pages:
            import pypdfium2 as pdfium

            pdf = pdfium.PdfDocument(args.file_path)
            print(f"PDF: 전체 {len(pdf)}페이지. --pages 옵션으로 범위를 지정하세요.")
            print_next_step(
                "범위 지정 후 필요한 페이지만 읽으세요.",
                "uv run --project skills/datasheet-intelligence "
                f'skills/datasheet-intelligence/scripts/read.py "{args.file_path}" --pages 1-5',
            )
            return
        read_pdf(args.file_path, _parse_page_range(args.pages))

    elif suffix == ".docx":
        read_docx(args.file_path)

    elif suffix in (".xlsx", ".xlsm"):
        read_xlsx(args.file_path)

    else:
        print(
            f"Error: 미지원 포맷 '{suffix}'. PDF, DOCX, XLSX를 지원합니다.",
            file=sys.stderr,
        )
        print_next_step(
            "지원 포맷 파일로 다시 실행하거나, docling structured 모드를 검토하세요.",
            "uv run --project skills/datasheet-intelligence --with docling "
            f'skills/datasheet-intelligence/scripts/read.py "{args.file_path}" --structured',
        )
        sys.exit(1)


if __name__ == "__main__":
    main()
