#!/usr/bin/env python3
"""파일 포맷을 자동 감지하여 키워드를 검색한다.

지원 포맷:
  - PDF/DOCX/XLSX
  - 기본: 텍스트/셀 단위 단순 검색 (빠름)
  - 구조적: docling 파싱 후 문맥(Heading, Paragraph) 인지 검색 (느림, 정확)
"""

import argparse
import json
import re
import sys
from pathlib import Path


def print_next_step(tip: str, command: str | None = None) -> None:
    """에러 직후 다음 액션을 동일 포맷(Tip/Try)으로 안내한다."""
    print(f"Tip: {tip}", file=sys.stderr)
    if command:
        print(f"Try: {command}", file=sys.stderr)


# ---------------------------------------------------------------------------
# 공통 검색 로직
# ---------------------------------------------------------------------------


def compile_patterns(queries: list[str], is_regex: bool) -> list[re.Pattern]:
    patterns = []
    for q in queries:
        flags = re.IGNORECASE
        patterns.append(re.compile(q if is_regex else re.escape(q), flags))
    return patterns


# ---------------------------------------------------------------------------
# 기본 검색 모드 (빠름)
# ---------------------------------------------------------------------------


def search_pdf_fast(
    path: Path,
    patterns: list[re.Pattern],
    context_chars: int,
    max_hits: int,
    unique_pages: bool = False,
) -> int:
    import pypdfium2 as pdfium

    pdf = pdfium.PdfDocument(path)
    hits = 0
    seen_keyword_pages: set[tuple[str, int]] = set()

    for page_idx, page in enumerate(pdf):
        if hits >= max_hits:
            break
        tp = page.get_textpage()
        text = tp.get_text_range()

        for pattern in patterns:
            for match in pattern.finditer(text):
                if hits >= max_hits:
                    break
                page_num = page_idx + 1
                dedupe_key = (pattern.pattern, page_num)
                if unique_pages and dedupe_key in seen_keyword_pages:
                    continue

                s, e = match.span()
                ctx = text[
                    max(0, s - context_chars) : min(len(text), e + context_chars)
                ].replace("\n", " ")
                print(
                    json.dumps(
                        {
                            "keyword": match.group(),
                            "page": page_num,
                            "context": f"...{ctx}...",
                            "mode": "fast",
                        },
                        ensure_ascii=False,
                    )
                )
                if unique_pages:
                    seen_keyword_pages.add(dedupe_key)
                hits += 1

        tp.close()
        page.close()
    return hits


def search_docx_fast(path: Path, patterns: list[re.Pattern], max_hits: int) -> int:
    from docx import Document

    doc = Document(path)
    hits = 0

    for i, para in enumerate(doc.paragraphs, 1):
        if hits >= max_hits:
            break
        text = para.text.strip()
        if not text:
            continue
        for pattern in patterns:
            if pattern.search(text):
                print(
                    json.dumps(
                        {
                            "keyword": pattern.pattern,
                            "paragraph": i,
                            "style": para.style.name if para.style else "Normal",
                            "text": text,
                            "mode": "fast",
                        },
                        ensure_ascii=False,
                    )
                )
                hits += 1
                if hits >= max_hits:
                    break
    return hits


def search_xlsx_fast(path: Path, patterns: list[re.Pattern], max_hits: int) -> int:
    from openpyxl import load_workbook

    wb = load_workbook(path, read_only=True, data_only=True)
    hits = 0

    for name in wb.sheetnames:
        if hits >= max_hits:
            break
        for row_idx, row in enumerate(wb[name].iter_rows(values_only=True), 1):
            if hits >= max_hits:
                break
            cells = [str(c) if c is not None else "" for c in row]
            row_text = " ".join(cells)
            for pattern in patterns:
                if pattern.search(row_text):
                    print(
                        json.dumps(
                            {
                                "keyword": pattern.pattern,
                                "sheet": name,
                                "row": row_idx,
                                "data": cells,
                                "mode": "fast",
                            },
                            ensure_ascii=False,
                        )
                    )
                    hits += 1
                    break
    wb.close()
    return hits


# ---------------------------------------------------------------------------
# 구조적 검색 모드 (docling)
# ---------------------------------------------------------------------------


def search_structured(
    path: Path, patterns: list[re.Pattern], max_hits: int, unique_pages: bool = False
) -> int:
    try:
        from docling.document_converter import DocumentConverter
        from docling.datamodel.document import TextItem, TableItem, SectionHeaderItem
    except ImportError:
        print("Error: --structured 모드는 docling이 필요합니다.", file=sys.stderr)
        print_next_step(
            "docling extra를 포함해 structured 검색으로 다시 실행하세요.",
            "uv run --project skills/datasheet-intelligence --with docling "
            "skills/datasheet-intelligence/scripts/search.py <file> <query> --structured",
        )
        sys.exit(1)

    print(f"[structured] {path.name} 파싱 및 검색 중...", file=sys.stderr)

    converter = DocumentConverter()
    try:
        result = converter.convert(path)
        doc = result.document
    except Exception as e:
        print(f"Error: docling 파싱 실패 — {e}", file=sys.stderr)
        print_next_step(
            "우선 fast 모드로 검색해 페이지 후보를 좁힌 뒤 structured를 재시도하세요.",
            "uv run --project skills/datasheet-intelligence "
            f'skills/datasheet-intelligence/scripts/search.py "{path}" <query> --unique-pages',
        )
        return 0

    hits = 0
    seen_keyword_pages: set[tuple[str, int]] = set()

    # 문서 순회
    # docling의 iterate_items는 (item, level)을 반환한다.
    # 검색 문맥 보존을 위해 현재 섹션 헤더를 추적한다.

    current_section = "Root"

    for item, level in doc.iterate_items():
        if hits >= max_hits:
            break

        if isinstance(item, SectionHeaderItem):
            current_section = item.text.strip()
            # 헤더 자체도 검색 대상이 될 수 있다.
            text = item.text
        elif isinstance(item, TextItem):
            text = item.text
        elif isinstance(item, TableItem):
            # 테이블은 dataframe 또는 HTML로 변환해 검색해야 한다.
            # 여기서는 dataframe 문자열 변환으로 단순화한다.
            try:
                text = item.export_to_dataframe().to_string()
            except:
                continue
        else:
            continue

        # 검색 수행
        for pattern in patterns:
            for match in pattern.finditer(text):
                if hits >= max_hits:
                    break

                # 문맥 정보 구성
                # TextItem, TableItem 등은 prov(페이지 정보)를 가질 수 있다.
                page_no = item.prov[0].page_no if item.prov else None
                if unique_pages and page_no is not None:
                    dedupe_key = (pattern.pattern, page_no)
                    if dedupe_key in seen_keyword_pages:
                        continue
                    seen_keyword_pages.add(dedupe_key)

                # 결과 출력
                print(
                    json.dumps(
                        {
                            "keyword": match.group(),
                            "section": current_section,
                            "page": page_no,
                            "type": type(item).__name__,
                            "content_snippet": (
                                text[:300].replace("\n", " ") + "..."
                                if len(text) > 300
                                else text.replace("\n", " ")
                            ),
                            "mode": "structured",
                        },
                        ensure_ascii=False,
                    )
                )

                hits += 1
            if hits >= max_hits:
                break

    return hits


# ---------------------------------------------------------------------------
# 메인 진입점
# ---------------------------------------------------------------------------


def main():
    parser = argparse.ArgumentParser(
        description="파일에서 키워드를 검색한다 (PDF/DOCX/XLSX)."
    )
    parser.add_argument("file_path", type=Path, help="대상 파일 경로")
    parser.add_argument("queries", nargs="+", help="검색 키워드 또는 정규식 패턴")
    parser.add_argument("--regex", action="store_true", help="쿼리를 정규식으로 처리")
    parser.add_argument(
        "--context-chars",
        type=int,
        default=200,
        help="PDF 매칭 주변 문자 수 (기본: 200)",
    )
    parser.add_argument(
        "--max-hits", type=int, default=20, help="최대 결과 수 (기본: 20)"
    )
    parser.add_argument(
        "--unique-pages",
        action="store_true",
        help="PDF/structured 검색에서 같은 페이지의 중복 결과를 키워드별로 1개로 축약",
    )
    parser.add_argument(
        "--structured", action="store_true", help="docling으로 레이아웃/구조 인식 검색"
    )
    args = parser.parse_args()

    if not args.file_path.exists():
        print(f"Error: 파일 없음 — {args.file_path}", file=sys.stderr)
        print_next_step(
            "파일 경로를 다시 확인하세요.",
            f'ls -l "{args.file_path}"',
        )
        sys.exit(1)

    patterns = compile_patterns(args.queries, args.regex)
    hits = 0

    if args.structured:
        hits = search_structured(
            args.file_path, patterns, args.max_hits, args.unique_pages
        )
    else:
        suffix = args.file_path.suffix.lower()
        if suffix == ".pdf":
            hits = search_pdf_fast(
                args.file_path,
                patterns,
                args.context_chars,
                args.max_hits,
                args.unique_pages,
            )
        elif suffix == ".docx":
            hits = search_docx_fast(args.file_path, patterns, args.max_hits)
        elif suffix in (".xlsx", ".xlsm"):
            hits = search_xlsx_fast(args.file_path, patterns, args.max_hits)
        else:
            print(f"Error: 미지원 포맷 '{suffix}'. (Fast 모드)", file=sys.stderr)
            print_next_step(
                "지원 포맷(PDF/DOCX/XLSX)으로 실행하거나 structured 모드로 전환하세요.",
                "uv run --project skills/datasheet-intelligence --with docling "
                f'skills/datasheet-intelligence/scripts/search.py "{args.file_path}" <query> --structured',
            )
            sys.exit(1)

    if hits == 0:
        print("결과 없음.", file=sys.stderr)
        print_next_step(
            "키워드를 확장하거나 --regex/--structured 옵션을 검토하세요.",
            "uv run --project skills/datasheet-intelligence "
            f'skills/datasheet-intelligence/scripts/search.py "{args.file_path}" <query1> <query2> --unique-pages',
        )


if __name__ == "__main__":
    main()
